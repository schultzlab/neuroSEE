#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# =============================================================================
# ** == STILL TO DO
# =============================================================================

# Python Imports
from datetime import datetime
import h5py
import numpy as np
import glob
import os

# Reading the Experiment Data sheet
from docx import Document

# NWB imports
from pynwb import NWBFile, get_manager
from pynwb import TimeSeries
from pynwb import NWBHDF5IO

# my scripts
import from_ini # this has crashed for some reason?????

# =============================================================================
# =============================================================================
# # CHECKS
# =============================================================================
# =============================================================================
# Create an output folder if it doesn't exist
# =============================================================================
# make a folder for mouse if it's not there
OUTPUT_DIR = "../nwb_scripts/mouse_nwb/"
if not os.path.exists(OUTPUT_DIR):
    os.makedirs(OUTPUT_DIR)

# =============================================================================
# Checking whether data you are trying to upload is already in the system
# Loop through the matlab files you are adding and check whether a mouse and
# file already exist within the mouse NWB folder
# =============================================================================
DATA_DIRECTORY = '/Users/solomia/Dropbox/2p_da/data_files/mat_files'
FILE_STRUCTURE = '/Users/solomia/Dropbox/2p_da/data_files/2018.03.07'
# =============================================================================
# SETTING DIRECTORIES 
# ** combine these into a function later- dictionary to hold folder names
# ** need to add date on the EXP DATA SHEET as iterations occur
# =============================================================================
# 2photon data, if folder is there set two photon dir 
# =============================================================================
if not os.path.exists(FILE_STRUCTURE + '/2P'):
    print('No 2P folder found')
else:
    TWOP_DIR = FILE_STRUCTURE + '/2P'
# =============================================================================    
# behaviour data, if folder is there set IR dir
# =============================================================================
if not os.path.exists(FILE_STRUCTURE + '/IR tracking'):
    print('No IR folder found')
else:
    IR_DIR = FILE_STRUCTURE + '/IR tracking'
# =============================================================================
# exp data sheet, if file is there set directory
# =============================================================================
if not os.path.exists(FILE_STRUCTURE + '/Formatted Experiment Data Sheet - 2018.03.07.docx'):
    print('No Data Sheet file found')
else:
    DATA_SHEET_DIR = FILE_STRUCTURE + '/Formatted Experiment Data Sheet - 2018.03.07.docx'

# =============================================================================
# READING EXP DATA SHEET, using this info to link to 2P and IR info
# =============================================================================
document = Document(DATA_SHEET_DIR)

# Prints out the headers (there aren't any anymore, all in the tables)
for i in range(0,len(document.tables)):
    print(document.tables[i].rows[0].cells[0].text)
    
# prints out data from tables
for table in document.tables:
    for row in table.rows:
        for cell in row.cells:
            print(cell.text)

# this holds the info for all the imaging
for i in range(0,len(document.tables)):
    if document.tables[i].rows[0].cells[0].text == 'Experiment details':
        print(document.tables[i].rows[0].cells[0].text)

import pandas as pd
# Dictionary with mouse info
subj_details = {}

# this gives first row of table with different mouse details
for i in range(0,len(document.tables)):
    if document.tables[i].rows[0].cells[0].text == 'Subject details':
#        Gets number and id
        for row in range(1,2):
            for cell in range(0,len(document.tables[2].rows[row].cells)-1):
                subj_details[document.tables[i].rows[row].cells[cell].text] = \
                document.tables[i].rows[row+1].cells[cell].text
# THIS KEEPS SAYING TUPLE INDEX OUT OF RANGE??????
    
# pandas data frames from the tables holding the data subject details
subj_pd = pd.DataFrame.from_dict(subj_details,orient='index')
# IR + 2P info
# =============================================================================
# Read existing NWB file and check for acquisitions
# Better to get info about data that is being attempted to be added
# then checking if it is there
# =============================================================================
io = NWBHDF5IO("conference_mouse.nwb")
nwbfile = io.read()

# loop through existing acquisitions to check if timeseries is already there
ts = nwbfile.get_acquisition('CaTimeSeriesRed_day2')

# =============================================================================
# Creating the NWB file
# =============================================================================
nwb_file = NWBFile('NWBFile Two-photon GCamp Imaging', 
            'testing', 
            '001', 
            file_create_date = datetime.now(tz=None),
            experimenter='Dr Ann Go',
            lab='Neural Coding',
            institution='Imperial College London',
            experiment_description='GCamp Imaging',
            session_id='test_001',
            nwb_version = '2.0',
            session_start_time = from_ini.session_start_time)

# =============================================================================
# Loading the matlab file
# =============================================================================
### find the matlab variable name inside the matlab file
def find_red(timeseries_R):
 """ Find first object with CaTimeSeriesRed anywhere in the name """
 if 'CaTimeSeriesRed' in timeseries_R:
     return timeseries_R

def find_green(timeseries_G):
 """ Find first object with CaTimeSeriesGreen anywhere in the name """
 if 'CaTimeSeriesGreen' in timeseries_G:
     return timeseries_G

# =============================================================================
# Creating a loop through each matlab file- extracting R and G
#     THIS WORKS!!!!
# =============================================================================
def get_mat_data(mat_data_directory):
    # make list of the mat files
    file_list = glob.glob(mat_data_directory + "/*.mat")
    for i in file_list:
        f = h5py.File(i,'r')
        data_R = f.get(f.visit(find_red))
        data_R = np.array(data_R)
        data_G = f.get(f.visit(find_green))
        data_G = np.array(data_G)
        time_series_R = TimeSeries(f.visit(find_red),
#                          'comments about the timeseries',
                            'description of timeseries',
                          data = data_R,
                          unit= 'Hz (i think)',
#                         need to get starting time or timestamps from somewhere
                          starting_time = 19.01,
                          rate = float(from_ini.rate)
                          )
        print(time_series_R)
        nwb_file.add_acquisition(time_series_R)

        time_series_G = TimeSeries(f.visit(find_green),
#                          'comments about the timeseries',
                          'description of timeseries',
                          data = data_G,
                          unit= 'Hz (i think)',
#                         need to get starting time or timestamps from somewhere
                          starting_time = 19.01,
                          rate = float(from_ini.rate)
                          )
        nwb_file.add_acquisition(time_series_G)

# =============================================================================
# =============================================================================
# # CHECKS (FROM SSC SCRIPT) not mine
# =============================================================================
# =============================================================================
# checking for an object in the file 
# this will need to check that there is a timeseries file in the mat file
def check_entry(file_name,obj):
    try:
        return file_name[obj]
    except KeyError:
        print (str(obj) +" does not exist")
        return []

# =============================================================================
# Calling the functions and building the files        
# =============================================================================

get_mat_data(DATA_DIRECTORY)

file = NWBHDF5IO(OUTPUT_DIR + "conference_mouse.nwb", manager=get_manager(), mode='w')
file.write(nwb_file)
file.close()



